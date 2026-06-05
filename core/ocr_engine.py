"""Motor OCR hibrido para PDFs.

La estrategia privilegia precision y trazabilidad:
  - conserva texto nativo cuando la pagina ya tiene una capa confiable;
  - renderiza y reconoce escaneos con Tesseract integrado en PyMuPDF;
  - compara una pasada original y otra mejorada para rescatar escaneos debiles;
  - prueba rotaciones adicionales cuando la lectura inicial parece pobre;
  - exporta una transcripcion editable en Word y/o TXT.

Los modelos de idioma viven en ``assets/tessdata`` para que el OCR funcione
tambien dentro del ejecutable portable generado por PyInstaller.
"""
from __future__ import annotations

from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Callable, Iterable, List, Literal, Optional
import os
import re
import sys
import unicodedata

import fitz
import numpy as np
from PIL import Image, ImageFilter, ImageOps

from .output_naming import output_stem_for_source


PrecisionMode = Literal["maximum", "balanced", "fast"]
OutputMode = Literal["docx_txt", "docx", "txt"]

_LANGUAGE_LABELS = {
    "spa": "Español",
    "eng": "Inglés",
}

_COMMON_WORDS = {
    "a", "al", "con", "de", "del", "el", "en", "es", "esta", "este",
    "la", "las", "los", "no", "para", "por", "que", "se", "su", "un",
    "una", "y", "the", "of", "to", "and", "in", "for", "is", "on", "with",
}


@dataclass
class OcrConfig:
    """Configuracion de reconocimiento y exportacion."""

    languages: str = "spa+eng"
    dpi: int = 300
    precision_mode: PrecisionMode = "maximum"
    preserve_native_text: bool = True
    enhance_scans: bool = True
    recover_rotated_pages: bool = True
    output_mode: OutputMode = "docx_txt"

    @property
    def export_docx(self) -> bool:
        return self.output_mode in ("docx_txt", "docx")

    @property
    def export_txt(self) -> bool:
        return self.output_mode in ("docx_txt", "txt")


@dataclass
class OcrJob:
    pdf_path: str
    output_dir: str
    base_name: str = ""
    tool_suffix: str = "OCR"
    add_tool_suffix: bool = True


@dataclass
class OcrPageResult:
    page_index: int
    text: str = ""
    paragraphs: List[str] = field(default_factory=list)
    method: str = "ocr"
    variant: str = ""
    quality_score: float = 0.0
    image_coverage: float = 0.0
    warning: str = ""

    @property
    def char_count(self) -> int:
        return len(re.sub(r"\s+", "", self.text))

    @property
    def word_count(self) -> int:
        return len(re.findall(r"\b[\w'-]+\b", self.text, flags=re.UNICODE))


@dataclass
class OcrJobResult:
    job: OcrJob
    output_path: str = ""
    docx_path: str = ""
    txt_path: str = ""
    page_results: List[OcrPageResult] = field(default_factory=list)
    success: bool = True
    cancelled: bool = False
    error: str = ""

    @property
    def native_pages(self) -> int:
        return sum(1 for page in self.page_results if page.method == "native")

    @property
    def ocr_pages(self) -> int:
        return sum(1 for page in self.page_results if page.method == "ocr")

    @property
    def warning_pages(self) -> int:
        return sum(1 for page in self.page_results if page.warning)

    @property
    def word_count(self) -> int:
        return sum(page.word_count for page in self.page_results)

    @property
    def text(self) -> str:
        return "\n\n".join(page.text for page in self.page_results if page.text)

    @property
    def average_quality(self) -> float:
        readable = [page.quality_score for page in self.page_results if page.text]
        return sum(readable) / len(readable) if readable else 0.0


@dataclass
class _ExtractionCandidate:
    text: str
    paragraphs: List[str]
    score: float
    variant: str
    rotation: int = 0


def get_tessdata_dir() -> Path:
    """Retorna la carpeta con modelos OCR.

    Prioridad:
      1. Variable de entorno PDFLEX_TESSDATA (override manual / testing)
      2. PyInstaller --onefile: sys._MEIPASS (directorio de extracción temporal)
      3. Nuitka standalone / PyInstaller --onedir: directorio del ejecutable
      4. Desarrollo: raíz del proyecto
    """
    override = os.environ.get("PDFLEX_TESSDATA", "").strip()
    candidates: list[Path] = []

    if override:
        candidates.append(Path(override))

    # PyInstaller --onefile
    meipass = getattr(sys, "_MEIPASS", None)
    if meipass:
        candidates.append(Path(meipass) / "assets" / "tessdata")
    elif getattr(sys, "frozen", False):
        # Nuitka standalone o PyInstaller --onedir
        candidates.append(Path(sys.executable).parent / "assets" / "tessdata")

    # Fallback de desarrollo (siempre al final)
    candidates.append(Path(__file__).resolve().parents[1] / "assets" / "tessdata")

    for candidate in candidates:
        if candidate.is_dir():
            return candidate
    return candidates[-1]


def available_languages() -> List[str]:
    """Lista de modelos OCR disponibles en la instalacion actual."""

    tessdata = get_tessdata_dir()
    if not tessdata.is_dir():
        return []
    return sorted(path.stem for path in tessdata.glob("*.traineddata"))


def describe_languages(languages: str) -> str:
    codes = [code for code in languages.split("+") if code]
    return " + ".join(_LANGUAGE_LABELS.get(code, code) for code in codes)


def validate_tessdata(languages: str) -> Optional[str]:
    """Valida que los modelos requeridos esten disponibles."""

    tessdata = get_tessdata_dir()
    if not tessdata.is_dir():
        return (
            "No se encontro la carpeta de modelos OCR. "
            "Reinstala PDFlex o verifica assets/tessdata."
        )

    missing = [
        code for code in languages.split("+")
        if code and not (tessdata / f"{code}.traineddata").is_file()
    ]
    if missing:
        return (
            "Faltan modelos OCR para: "
            + ", ".join(_LANGUAGE_LABELS.get(code, code) for code in missing)
            + "."
        )
    return None


class OcrEngine:
    """Extrae texto nativo o aplica OCR por pagina y genera entregables."""

    def run_batch(
        self,
        jobs: List[OcrJob],
        config: OcrConfig,
        progress: Optional[Callable[[int, int, str], None]] = None,
        should_cancel: Optional[Callable[[], bool]] = None,
        result_ready: Optional[Callable[[OcrJobResult], None]] = None,
    ) -> List[OcrJobResult]:
        model_error = validate_tessdata(config.languages)
        if model_error:
            raise RuntimeError(model_error)

        total_pages = self._count_pages(jobs)
        completed_pages = 0
        results: List[OcrJobResult] = []

        for job_index, job in enumerate(jobs):
            if _is_cancelled(should_cancel):
                break

            result = self.run_job(
                job,
                config,
                progress=lambda local_done, local_total, msg, base=completed_pages: (
                    progress(base + local_done, total_pages, msg) if progress else None
                ),
                should_cancel=should_cancel,
            )
            results.append(result)
            if result_ready:
                result_ready(result)
            completed_pages += len(result.page_results)

            if result.cancelled:
                break

            if progress:
                progress(
                    completed_pages,
                    total_pages,
                    f"Documento {job_index + 1}/{len(jobs)} completado",
                )

        if progress and not _is_cancelled(should_cancel):
            progress(total_pages, total_pages, "Extracción completada")

        return results

    def run_job(
        self,
        job: OcrJob,
        config: OcrConfig,
        progress: Optional[Callable[[int, int, str], None]] = None,
        should_cancel: Optional[Callable[[], bool]] = None,
    ) -> OcrJobResult:
        try:
            doc = fitz.open(job.pdf_path)
        except Exception as exc:
            return OcrJobResult(job=job, success=False, error=str(exc))

        pages: List[OcrPageResult] = []
        total = doc.page_count
        try:
            for page_index in range(total):
                if _is_cancelled(should_cancel):
                    return OcrJobResult(
                        job=job,
                        page_results=pages,
                        success=False,
                        cancelled=True,
                        error="Proceso cancelado por el usuario.",
                    )

                if progress:
                    progress(
                        page_index,
                        total,
                        f"{Path(job.pdf_path).name}: analizando página {page_index + 1}/{total}",
                    )

                page = doc[page_index]
                pages.append(
                    self._extract_page(page, page_index, config, should_cancel)
                )

                if progress:
                    progress(
                        page_index + 1,
                        total,
                        f"{Path(job.pdf_path).name}: página {page_index + 1}/{total} lista",
                    )
        except _CancelledError:
            return OcrJobResult(
                job=job,
                page_results=pages,
                success=False,
                cancelled=True,
                error="Proceso cancelado por el usuario.",
            )
        except Exception as exc:
            return OcrJobResult(
                job=job,
                page_results=pages,
                success=False,
                error=str(exc),
            )
        finally:
            doc.close()

        try:
            docx_path, txt_path = self._export(job, pages, config)
        except Exception as exc:
            return OcrJobResult(
                job=job,
                page_results=pages,
                success=False,
                error=f"No se pudo generar el entregable: {exc}",
            )

        primary = docx_path or txt_path
        return OcrJobResult(
            job=job,
            output_path=primary,
            docx_path=docx_path,
            txt_path=txt_path,
            page_results=pages,
            success=bool(primary),
            error="" if primary else "No se generó ningún entregable.",
        )

    @staticmethod
    def _count_pages(jobs: Iterable[OcrJob]) -> int:
        total = 0
        for job in jobs:
            try:
                with fitz.open(job.pdf_path) as doc:
                    total += doc.page_count
            except Exception:
                total += 1
        return max(1, total)

    def _extract_page(
        self,
        page: fitz.Page,
        page_index: int,
        config: OcrConfig,
        should_cancel: Optional[Callable[[], bool]],
    ) -> OcrPageResult:
        native_paragraphs = _extract_blocks(page)
        native_text = _join_paragraphs(native_paragraphs)
        image_coverage = _page_image_coverage(page)

        if (
            config.preserve_native_text
            and _is_reliable_native_text(native_text, image_coverage)
        ):
            return OcrPageResult(
                page_index=page_index,
                text=native_text,
                paragraphs=native_paragraphs,
                method="native",
                variant="Capa de texto original",
                quality_score=1.0,
                image_coverage=image_coverage,
            )

        if _is_cancelled(should_cancel):
            raise _CancelledError

        sideways_hint = (
            config.recover_rotated_pages
            and config.precision_mode != "fast"
            and _page_looks_sideways(page)
        )
        pix = page.get_pixmap(dpi=config.dpi, colorspace=fitz.csRGB, alpha=False)
        image: Optional[Image.Image] = None

        if sideways_hint:
            image = _pixmap_to_pil(pix)
            best: Optional[_ExtractionCandidate] = None
            for rotation in (90, 270):
                if _is_cancelled(should_cancel):
                    raise _CancelledError
                candidate = self._ocr_image(image, config, rotation=rotation, enhanced=False)
                best = candidate if best is None else _best_candidate(best, candidate)
                if _is_high_confidence(best.text, best.score):
                    break
            if best is None or not _is_high_confidence(best.text, best.score):
                best = _best_candidate(
                    best,
                    self._ocr_pixmap(
                        pix,
                        config,
                        variant="OCR / imagen original",
                    ),
                ) if best else self._ocr_pixmap(
                    pix,
                    config,
                    variant="OCR / imagen original",
                )
        else:
            best = self._ocr_pixmap(
                pix,
                config,
                variant="OCR / imagen original",
            )

        if (
            config.enhance_scans
            and config.precision_mode != "fast"
            and _needs_enhancement(best.text, best.score)
        ):
            if _is_cancelled(should_cancel):
                raise _CancelledError
            if image is None:
                image = _pixmap_to_pil(pix)
            best = _best_candidate(
                best,
                self._ocr_image(
                    image,
                    config,
                    rotation=best.rotation,
                    enhanced=True,
                ),
            )

        should_try_rotation = (
            not sideways_hint
            and config.recover_rotated_pages
            and config.precision_mode != "fast"
            and _needs_rotation_recovery(best.text, best.score)
        )
        if should_try_rotation:
            if image is None:
                image = _pixmap_to_pil(pix)
            for rotation in (90, 270, 180):
                if _is_cancelled(should_cancel):
                    raise _CancelledError
                best = _best_candidate(
                    best,
                    self._ocr_image(image, config, rotation=rotation, enhanced=False),
                )
                if _is_high_confidence(best.text, best.score):
                    break

        warning = ""
        if not best.text.strip() and native_text.strip():
            best = _ExtractionCandidate(
                text=native_text,
                paragraphs=native_paragraphs,
                score=_score_text(native_text),
                variant="Capa original de respaldo",
            )
            warning = "El OCR no encontró texto; se conservó la capa original."
        elif not best.text.strip():
            warning = "No se detectó texto legible en esta página."
        elif best.score < 0.48:
            warning = "La página tiene legibilidad baja; conviene revisarla."

        return OcrPageResult(
            page_index=page_index,
            text=best.text,
            paragraphs=best.paragraphs,
            method="ocr",
            variant=best.variant,
            quality_score=best.score,
            image_coverage=image_coverage,
            warning=warning,
        )

    def _ocr_image(
        self,
        image: Image.Image,
        config: OcrConfig,
        *,
        rotation: int,
        enhanced: bool,
    ) -> _ExtractionCandidate:
        prepared = image
        if rotation:
            prepared = prepared.rotate(rotation, expand=True, fillcolor="white")
        if enhanced:
            prepared = _enhance_image(prepared)

        if prepared.mode != "RGB":
            prepared = prepared.convert("RGB")
        pix = fitz.Pixmap(
            fitz.csRGB,
            prepared.width,
            prepared.height,
            prepared.tobytes(),
            False,
        )
        parts = ["OCR"]
        if enhanced:
            parts.append("mejora de contraste")
        if rotation:
            parts.append(f"rotación {rotation} grados")
        return self._ocr_pixmap(
            pix,
            config,
            variant=" / ".join(parts),
            rotation=rotation,
        )

    def _ocr_pixmap(
        self,
        pix: fitz.Pixmap,
        config: OcrConfig,
        *,
        variant: str,
        rotation: int = 0,
    ) -> _ExtractionCandidate:
        ocr_pdf = pix.pdfocr_tobytes(
            language=config.languages,
            tessdata=str(get_tessdata_dir()),
            compress=True,
        )

        with fitz.open("pdf", ocr_pdf) as doc:
            paragraphs = _extract_blocks(doc[0])
        text = _join_paragraphs(paragraphs)

        return _ExtractionCandidate(
            text=text,
            paragraphs=paragraphs,
            score=_score_text(text),
            variant=variant,
            rotation=rotation,
        )

    def _export(
        self,
        job: OcrJob,
        pages: List[OcrPageResult],
        config: OcrConfig,
    ) -> tuple[str, str]:
        out_dir = Path(job.output_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        base = output_stem_for_source(
            job.base_name or job.pdf_path,
            tool_suffix=job.tool_suffix,
            add_tool_suffix=job.add_tool_suffix,
            fallback="documento",
        )

        txt_path = ""
        if config.export_txt:
            path = out_dir / f"{base}.txt"
            path.write_text(
                _build_txt_export(Path(job.pdf_path).name, pages),
                encoding="utf-8-sig",
            )
            txt_path = str(path)

        docx_path = ""
        if config.export_docx:
            path = out_dir / f"{base}.docx"
            _write_docx(path, Path(job.pdf_path).name, pages)
            docx_path = str(path)

        return docx_path, txt_path


class _CancelledError(Exception):
    pass


def ocr_job_result_to_dict(result: OcrJobResult) -> dict:
    """Serializa un resultado para comunicarlo desde el proceso OCR."""

    return asdict(result)


def ocr_job_result_from_dict(data: dict) -> OcrJobResult:
    """Reconstruye un resultado recibido desde el proceso OCR."""

    return OcrJobResult(
        job=OcrJob(**data["job"]),
        output_path=data.get("output_path", ""),
        docx_path=data.get("docx_path", ""),
        txt_path=data.get("txt_path", ""),
        page_results=[
            OcrPageResult(**page_data)
            for page_data in data.get("page_results", [])
        ],
        success=bool(data.get("success", False)),
        cancelled=bool(data.get("cancelled", False)),
        error=data.get("error", ""),
    )


def _is_cancelled(callback: Optional[Callable[[], bool]]) -> bool:
    return bool(callback and callback())


def _pixmap_to_pil(pix: fitz.Pixmap) -> Image.Image:
    """Crea una imagen PIL solo cuando una variante secundaria la necesita."""

    return Image.frombytes("RGB", (pix.width, pix.height), pix.samples)


def _page_looks_sideways(page: fitz.Page) -> bool:
    """Detecta texto lateral con una miniatura barata antes de ejecutar OCR."""

    try:
        pix = page.get_pixmap(dpi=72, colorspace=fitz.csGRAY, alpha=False)
        gray = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width)
        ink = gray < 210
        if int(ink.sum()) < 120:
            return False
        row_variation = float(ink.sum(axis=1).std())
        column_variation = float(ink.sum(axis=0).std())
        return column_variation > row_variation * 1.55
    except Exception:
        return False


def _best_candidate(
    current: _ExtractionCandidate,
    candidate: _ExtractionCandidate,
) -> _ExtractionCandidate:
    return candidate if candidate.score > current.score else current


def _extract_blocks(page: fitz.Page) -> List[str]:
    paragraphs: List[str] = []
    for block in page.get_text("blocks", sort=True):
        if len(block) < 5:
            continue
        text = _normalize_text(str(block[4]))
        if text:
            paragraphs.append(text)
    return paragraphs


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(line.rstrip() for line in text.splitlines())
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _join_paragraphs(paragraphs: Iterable[str]) -> str:
    return "\n\n".join(part.strip() for part in paragraphs if part.strip()).strip()


def _page_image_coverage(page: fitz.Page) -> float:
    page_area = max(1.0, page.rect.width * page.rect.height)
    image_area = 0.0
    try:
        for info in page.get_image_info():
            rect = fitz.Rect(info.get("bbox", (0, 0, 0, 0)))
            rect.intersect(page.rect)
            image_area += max(0.0, rect.width) * max(0.0, rect.height)
    except Exception:
        return 0.0
    return min(1.0, image_area / page_area)


def _is_reliable_native_text(text: str, image_coverage: float) -> bool:
    compact = re.sub(r"\s+", "", text)
    if not compact:
        return False
    score = _score_text(text)
    # Un escaneo puede traer solo encabezados o sellos como texto nativo.
    # Si una imagen ocupa casi toda la pagina, exigimos una capa mas completa.
    if image_coverage >= 0.72 and len(compact) < 350:
        return False
    if len(compact) >= 80 and score >= 0.58:
        return True
    return len(compact) >= 30 and image_coverage < 0.35 and score >= 0.62


def _score_text(text: str) -> float:
    compact = re.sub(r"\s+", "", text)
    if not compact:
        return 0.0

    valid_chars = sum(
        1 for char in compact
        if char.isalnum() or unicodedata.category(char)[0] in {"L", "N", "P", "S"}
    )
    valid_ratio = valid_chars / len(compact)

    tokens = re.findall(r"[^\W\d_][\w'-]*", text.casefold(), flags=re.UNICODE)
    if tokens:
        sane_tokens = sum(
            1 for token in tokens
            if (
                (len(token) >= 2 or token in {"a", "e", "o", "u", "y"})
                and len(token) <= 28
                and not re.search(r"(.)\1{4,}", token)
            )
        )
        sane_ratio = sane_tokens / len(tokens)
        common_hits = sum(1 for token in tokens if token in _COMMON_WORDS)
        common_score = min(1.0, common_hits / 5.0)
        average_token_length = sum(len(token) for token in tokens) / len(tokens)
        token_length_score = min(1.0, max(0.0, (average_token_length - 1.0) / 3.5))
    else:
        sane_ratio = 0.0
        common_score = 0.0
        token_length_score = 0.0

    length_score = min(1.0, len(compact) / 220.0)
    replacement_penalty = min(0.35, text.count("\ufffd") / max(1, len(compact)))
    score = (
        0.24 * valid_ratio
        + 0.24 * sane_ratio
        + 0.18 * length_score
        + 0.20 * common_score
        + 0.14 * token_length_score
        - replacement_penalty
    )
    return max(0.0, min(1.0, score))


def _needs_rotation_recovery(text: str, score: float) -> bool:
    compact = re.sub(r"\s+", "", text)
    tokens = re.findall(r"[^\W\d_][\w'-]*", text.casefold(), flags=re.UNICODE)
    fragmented = (
        len(tokens) >= 12
        and sum(1 for token in tokens if len(token) == 1) / len(tokens) > 0.38
    )
    return len(compact) < 45 or score < 0.52 or fragmented


def _needs_enhancement(text: str, score: float) -> bool:
    """Evita una segunda pasada cuando la lectura original ya es muy sólida."""

    return not _is_high_confidence(text, score)


def _is_high_confidence(text: str, score: float) -> bool:
    compact = re.sub(r"\s+", "", text)
    return len(compact) >= 45 and score >= 0.90 and not _needs_rotation_recovery(text, score)


def _enhance_image(image: Image.Image) -> Image.Image:
    """Realza contraste sin destruir trazos finos ni escalas de grises."""

    gray = ImageOps.grayscale(image)
    gray = ImageOps.autocontrast(gray, cutoff=1)
    gray = gray.filter(ImageFilter.UnsharpMask(radius=1.1, percent=150, threshold=3))
    return gray.convert("RGB")


def _build_txt_export(source_name: str, pages: List[OcrPageResult]) -> str:
    lines = [
        Path(source_name).stem,
        f"Fuente: {source_name}",
        f"Transcripción generada por PDFlex OCR el {datetime.now():%Y-%m-%d %H:%M}",
        "",
    ]
    for page in pages:
        lines.extend([
            f"==================== PÁGINA {page.page_index + 1} ====================",
            "",
            page.text.strip() or "[Sin texto legible]",
            "",
        ])
    return "\n".join(lines).rstrip() + "\n"


def _write_docx(path: Path, source_name: str, pages: List[OcrPageResult]) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_BREAK, WD_LINE_SPACING
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError(
            "La exportacion Word requiere python-docx. "
            "Instala dependencias con: pip install -r requirements.txt"
        ) from exc

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    title = doc.add_heading(Path(source_name).stem, level=0)
    title.paragraph_format.space_after = Pt(3)

    source = doc.add_paragraph()
    source.paragraph_format.space_after = Pt(10)
    run = source.add_run(f"Fuente: {source_name} | Transcripción PDFlex OCR")
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(100, 100, 100)

    for page_number, page in enumerate(pages, start=1):
        heading = doc.add_heading(f"Página {page_number}", level=2)
        heading.paragraph_format.space_before = Pt(3)
        heading.paragraph_format.space_after = Pt(5)

        if page.paragraphs:
            for paragraph_text in page.paragraphs:
                paragraph = doc.add_paragraph()
                lines = paragraph_text.splitlines() or [paragraph_text]
                for line_index, line in enumerate(lines):
                    if line_index:
                        paragraph.add_run().add_break(WD_BREAK.LINE)
                    paragraph.add_run(line)
        else:
            paragraph = doc.add_paragraph("[Sin texto legible]")
            paragraph.runs[0].italic = True

        if page_number < len(pages):
            doc.add_page_break()

    doc.core_properties.title = Path(source_name).stem
    doc.core_properties.subject = "Transcripción OCR editable"
    doc.core_properties.author = "PDFlex - GRUPO OCMX"
    doc.core_properties.comments = "Generado localmente con PDFlex OCR."
    doc.save(str(path))


_UNSAFE_CHARS = r'\/:*?"<>|'


def _sanitize_filename(name: str) -> str:
    for char in _UNSAFE_CHARS:
        name = name.replace(char, "_")
    return name.strip() or "documento"
